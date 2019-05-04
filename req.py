import sys, os, docx
from PyQt5 import QtCore, QtGui, QtWidgets
from req_ui import *
from docx.shared import Pt

with open("input.txt", encoding = "utf-8") as f:
    input = f.read().split("\n")

d = {}
d["ФИО"] = []
d["Адрес"] = []

for line in input:
    if line != "":
        title = line.split(": ",1)[0]
        if "Файл" in title and title != "Файл ИЦ":
            d["Файл"] = line.split(": ",1)[1]
        if title == "ФИО" or title == "Адрес":
            d[title].append(line.split(": ",1)[1])
        else:
            d[title] = line.split(": ",1)[1]

class MyWin(QtWidgets.QMainWindow):
    def __init__ (self, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)

        # Заполняем поля сохраненными значениями
        self.ui.dataEdit.setText(d["Дата"])
        self.ui.numberEdit.setText(d["Номер материала/дела"])
        self.ui.placeEdit.setText(d["Населенный пункт"])
        if int(d["Материал/Дело[0/1]"]) == 0:
            self.ui.mpRadioButton.setChecked(True)
        if int(d["Материал/Дело[0/1]"]) == 1:
            self.ui.udRadioButton.setChecked(True)
        self.changeFabula()

        self.ui.personInput.setPlainText("")
        for i in range(len(d["ФИО"])):
            self.ui.personInput.appendPlainText("ФИО: " + d["ФИО"][i])
            self.ui.personInput.appendPlainText("Адрес: " + d["Адрес"][i] + "\n")

        # Обработчик событий
        self.ui.mpRadioButton.clicked.connect(self.changeFabula)
        self.ui.udRadioButton.clicked.connect(self.changeFabula)
        self.ui.readyButton.clicked.connect(self.reqFunc)

    def changeFabula(self):
        if self.ui.mpRadioButton.isChecked():
            fab = "Фабула по материалу"
        if self.ui.udRadioButton.isChecked():
            fab = "Фабула по делу"
        self.ui.fabulaInput.setPlainText(d[fab])

    def reqFunc(self):
        self.ui.label_6.setText("Успешно")
        self.writeFunc()

        doc = docx.Document("Образцы" + os.sep + d["Файл"] + ".docx")

        for p in doc.paragraphs:
            for run in p.runs:
                if "COUNTRY" in run.text: # населенный пункт
                    run.text = run.text[:-len("COUNTRY")-1] + d["Населенный пункт"] + "»"

                if run.text == "DATE": # номер дела или материала
                    run.text = d["Дата"] + " " * 16 + d["Номер материала/дела"]

                if int(d["Материал/Дело[0/1]"]) == 0:
                    fab = "Фабула по материалу"
                else:
                    fab = "Фабула по делу"

                if "FABULA" in run.text:
                    run.text = d[fab] + run.text[len("FABULA"):]


                if run.text == "FIO": # имя
                    while len(p.runs) < len(d["ФИО"]) * 2: # добавляем runs
                        p.add_run()

                    for i in range(len(p.runs)):
                        if i % 2 == 0: # сначала имя
                            p.runs[i].text = str(i//2 + 1) + ") " + d["ФИО"][i//2]
                            p.runs[i].bold = True
                        else: # затем адрес
                            p.runs[i].text = ", " + d["Адрес"][i//2]
                            if i < len(p.runs)-1 :
                                p.runs[i].text += "\t\n\t"
                        p.runs[i].font.name = "Times New Roman"
                        p.runs[i].font.size = Pt(14)

        doc.save('Запросы готовые.docx')

        # требования ИЦ, ГИАЦ
        for count in range(len(d["ФИО"])):
            doc = docx.Document("Образцы" + os.sep + d["Файл ИЦ"] + ".docx") # файл ИЦ
            for p in doc.paragraphs:
                for run in p.runs:
                    if "FAMILIA" in run.text:
                        run.text = d["ФИО"][count].split()[0]
                    if "IO" in run.text:
                        run.text = d["ФИО"][count].split()[1] + " " + d["ФИО"][count].split()[2][:-1]
                    if "YEAR" in run.text:
                        run.text = d["ФИО"][count].split(",")[1]
                    if "ADRESS" in run.text:
                        run.text =  d["Адрес"][count].split(":")[1]
                    if "NUMBER" in run.text:
                        run.text =  "№" + d["Номер материала/дела"]
                    if "CURRENT" in run.text:
                        run.text =  " " * 6 + d["Дата"]

            doc.save('Требование ИЦ, ГИАЦ ' + str(count + 1) + " " + d["ФИО"][count].split()[0] + '.docx')


    def writeFunc(self):
        # Присваиваем значения полей в переменную input, котрую затем запишем в файл
        for line in input:
            if line != "":
                title = line.split(": ",1)[0]
                if title == "ФИО" or title == "Адрес":
                    input.remove(line)
                if title == "Дата":
                    input[input.index(line)] = title + ": " + self.ui.dataEdit.text()
                if title == "Материал/Дело[0/1]":
                    if self.ui.mpRadioButton.isChecked():
                        input[input.index(line)] = title + ": " + "0"
                    else:
                        input[input.index(line)] = title + ": " + "1"
                if title == "Номер материала/дела":
                    input[input.index(line)] = title + ": " + self.ui.numberEdit.text()

                if title == "Фабула по материалу" and self.ui.mpRadioButton.isChecked():
                    input[input.index(line)] = title + ": " + self.ui.fabulaInput.toPlainText()
                if title == "Фабула по делу" and self.ui.udRadioButton.isChecked():
                    input[input.index(line)] = title + ": " + self.ui.fabulaInput.toPlainText()

                if title == "Населенный пункт":
                    input[input.index(line)] = title + ": " + self.ui.placeEdit.text()

        # очищаем input от пустых элементов
        while "" in input:
            input.remove("")

        # вписываем в конец input содержимое поля personInput
        persons = self.ui.personInput.toPlainText().split("\n")
        for line in persons:
            if line != "":
                input.append(line)

        # добавляем в input пустые элементы через каждую строку, чтобы корректно срабатывало последующее считывание из файла
        tInput = input
        for i in range(len(tInput) * 2 - 1):
            if i % 2 != 0:
                input.insert(i, "")

        # отредактированную переменную записываем в файл
        with open("input.txt", "w") as f:
            f.write("\n".join(input))


if __name__ == "__main__":
	app = QtWidgets.QApplication(sys.argv)
	myapp = MyWin()
	myapp.show()
	sys.exit(app.exec_())